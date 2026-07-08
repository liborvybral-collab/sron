#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Generate Word document: 实荣与地矿二院合作契合点分析"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


def set_cell_shading(cell, color_hex):
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn("w:shd"), {
        qn("w:fill"): color_hex,
        qn("w:val"): "clear",
    })
    shading.append(shd)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "微软雅黑"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    return h


def add_para(doc, text, bold=False, size=11, align=None, space_after=6):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    run.bold = bold
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.35
    return p


def add_bullet(doc, text, size=11):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    p.paragraph_format.line_spacing = 1.35
    return p


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        set_cell_shading(hdr_cells[i], "1F4E79")
        for p in hdr_cells[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.name = "微软雅黑"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
                run.font.size = Pt(10)
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = val
            for p in cells[ci].paragraphs:
                for run in p.runs:
                    run.font.name = "宋体"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                    run.font.size = Pt(10)
        if ri % 2 == 1:
            for cell in cells:
                set_cell_shading(cell, "F2F7FB")
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table


def build_document():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("河南实荣筒仓工程有限公司\n与河南省第二地质矿产调查院有限公司\n合作契合点分析报告")
    tr.bold = True
    tr.font.size = Pt(18)
    tr.font.name = "微软雅黑"
    tr._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    tr.font.color.rgb = RGBColor(31, 78, 121)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("（商务交流参考文件）")
    sr.font.size = Pt(11)
    sr.font.name = "宋体"
    sr._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    sr.font.color.rgb = RGBColor(100, 100, 100)

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ir = info.add_run("编制单位：河南实荣筒仓工程有限公司（SRON）\n文件日期：2026年7月")
    ir.font.size = Pt(10)
    ir.font.name = "宋体"
    ir._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    ir.font.color.rgb = RGBColor(120, 120, 120)
    doc.add_paragraph()

    # Executive summary
    add_heading(doc, "一、总体判断", level=1)
    add_para(
        doc,
        "河南省第二地质矿产调查院有限公司（以下简称「地矿二院」）隶属豫地科技集团，"
        "已从传统地质勘查技术服务，转向「探、采、工、贸」一体化的境外矿产资源开发。"
        "其在坦桑尼亚、几内亚等地的金矿、石墨矿开发及产业园建设，对矿石储存、精矿缓存、"
        "封闭堆场、装卸转运等散料储运设施存在持续且刚性的需求。"
    )
    add_para(
        doc,
        "河南实荣筒仓工程有限公司（以下简称「实荣」）是国内领先的散装物料储运工程EPC总承包商，"
        "在矿粉钢板仓、封闭料场、海外装配式储运系统等领域具备成熟能力。"
        "双方契合度较高，合作空间主要集中在海外矿山储运配套、国内矿山封闭堆场及工程分包协作等方向。",
        bold=False,
    )

    add_table(
        doc,
        ["评估维度", "结论"],
        [
            ["合作可行性", "高——需求真实、能力互补、赛道不重叠"],
            ["最佳切入项目", "坦桑尼亚ML644石墨矿精矿库及缓存库"],
            ["合作模式", "贵院总包/主导，实荣储运系统分包或设备供货"],
            ["主要风险", "国企决策链较长，需从具体子项目逐步切入"],
        ],
        col_widths=[4, 12],
    )

    # Section 2 - About parties
    add_heading(doc, "二、双方基本情况", level=1)

    add_heading(doc, "2.1 地矿二院概况", level=2)
    add_bullet(doc, "隶属单位：河南省豫地科技集团有限公司（省属国企）")
    add_bullet(doc, "注册地址：河南省郑州市高新区莲花街56号")
    add_bullet(doc, "注册资本：5亿元人民币")
    add_bullet(doc, "主业：矿产资源勘查开发、矿山建设、地质灾害治理、生态修复、对外承包工程")
    add_bullet(doc, "海外布局：坦桑尼亚、几内亚等，拥有金、石墨等矿权40余处")
    add_bullet(doc, "重点项目：ML553金矿、ML644石墨矿、环维多利亚湖资源综合利用产业园")
    add_bullet(doc, "资质优势：各类资质29项（甲级12项），拥有坦桑尼亚CRB工程承包商资质")

    add_heading(doc, "2.2 实荣概况", level=2)
    add_bullet(doc, "企业定位：全球大宗散料储运工程系统解决方案提供商")
    add_bullet(doc, "核心能力：钢板仓工程、封闭料场（储道品牌）、码头装卸、智能化储运管控")
    add_bullet(doc, "服务链条：工程设计 → 装备制造 → 施工管理 → 安装调试 → 运维服务")
    add_bullet(doc, "海外业绩：40余个国家和地区，900余座钢板仓，储料总容量超300万吨")
    add_bullet(doc, "非洲经验：卢旺达装配式钢板仓等项目")
    add_bullet(doc, "资质认证：建筑工程/机电/钢结构总承包，ISO9001、SGS、TUV等")

    # Section 3 - Fit points
    add_heading(doc, "三、六大合作契合点", level=1)

    fits = [
        (
            "契合点一：坦桑尼亚ML644石墨矿项目（近期最具体）",
            "地矿二院2024年9月开工的ML644超大型石墨矿项目总投资约1亿美元，规划建设精矿脱水烘干车间、"
            "精矿分级及包装车间、精矿库及湿料缓存库（建筑面积约2381㎡）、药剂制备车间等。",
            [
                "石墨精粉、湿料/干料缓冲仓及钢板仓系统",
                "封闭料场/网架矿棚（防尘、防雨、便于装卸）",
                "精矿储存系统工艺设计 + 钢构制造 + 现场指导安装",
                "螺栓装配式仓方案：便于海运、现场快速组装，适合非洲工期紧的特点",
            ],
            "建议作为本次拜访重点推进的具体项目切口。",
        ),
        (
            "契合点二：坦桑尼亚金矿及环维多利亚湖产业园",
            "地矿二院在坦桑尼亚已建成ML553金矿（日处理矿石500吨级浮选厂），"
            "并推进环维多利亚湖资源综合利用产业园（商务部统计的境外经贸合作区），"
            "规划涵盖商务办公、化验测试、机械设备加工、金矿开发等分区。",
            [
                "金精粉/尾矿储存：矿粉钢板仓、封闭堆场",
                "产业园物流配套：散料中转仓、装卸系统",
                "园区设备加工区：装配式钢板仓标准化产品输出",
                "仓储管理：粮情监测技术可迁移为矿粉温湿度/粉尘监测",
            ],
            "地矿二院已明确将「仓储管理」纳入海外运营人才体系建设。",
        ),
        (
            "契合点三：几内亚等资源基地",
            "地矿二院在几内亚拥有金、石墨等矿权及固定生产基地，形成探采一体化产业布局。",
            [
                "矿山营地散料仓及矿石堆场封闭",
                "海外EPC分包合作（地矿二院有CRB资质，实荣提供储运系统）",
                "模块化钢构及装配式仓出口",
            ],
            "实荣海外工程经验可支撑几内亚基地建设。",
        ),
        (
            "契合点四：国内矿山与生态修复项目",
            "国内方面，地矿二院拥有嵩县金矿开发（合资嵩县山金矿业）及南太行山水林田湖草生态修复等项目。",
            [
                "矿山原矿/粉矿临时堆存封闭料场",
                "生态修复项目物料暂存仓",
                "与龙佰等合作伙伴形成「矿企+地勘+储运EPC」组合供货",
            ],
            "国内项目决策相对清晰，可作为合作试点。",
        ),
        (
            "契合点五：商业模式合作（中长期）",
            "双方可在商业模式层面探索多种合作路径。",
            [
                "工程分包：地矿二院总包矿山/选厂，实荣分包储运系统",
                "设备供货+指导安装：国内制造、海外组装，降本增效",
                "联合出海：地矿二院拿矿权/建设权，实荣提供储运标准与制造能力",
                "进出口通道：联合开展非洲备件、钢结构出口",
                "技术咨询：参与海外项目储运方案可研（前期投入小、易切入）",
            ],
            "建议从单一子项目技术咨询或方案比选起步。",
        ),
        (
            "契合点六：能力与资质互补",
            "双方在资质、能力、市场定位上形成互补，非竞争关系。",
            [
                "实荣：矿粉/矿石钢板仓、封闭网架料场、海外EPC、装配式螺栓仓",
                "地矿二院：矿权获取、地质勘查、矿山建设总包、海外基地运营",
                "联合优势：勘查开发+储运配套一体化解决方案",
            ],
            "叙事定位：产业链上下游协作，而非同业竞争。",
        ),
    ]

    for i, (title, background, items, note) in enumerate(fits, 1):
        add_heading(doc, f"3.{i} {title}", level=2)
        add_para(doc, background)
        add_para(doc, "实荣可对接：", bold=True)
        for item in items:
            add_bullet(doc, item)
        add_para(doc, f"▶ {note}", size=10)

    # Section 4 - Capability mapping
    add_heading(doc, "四、能力匹配对照表", level=1)
    add_table(
        doc,
        ["地矿二院需求场景", "实荣对应方案", "推荐产品/技术"],
        [
            ["石墨精矿储存", "精矿粉钢板仓系统", "焊接仓/装配式仓"],
            ["湿料/干料缓存", "缓冲仓+输送系统", "矿粉仓+配套钢构"],
            ["露天矿石堆场", "封闭网架料场", "储道品牌网架矿棚"],
            ["金矿精矿储运", "精矿仓+装卸系统", "矿粉钢板仓+装车系统"],
            ["海外工期紧张", "模块化快速建仓", "螺栓装配式钢板仓"],
            ["园区物流中转", "散料中转仓", "中转仓+电气集成"],
            ["粉尘/潮气控制", "密闭储存+监测", "密闭仓+温湿度监测"],
            ["国内矿山堆存", "临时堆场封闭", "网架料场/膜结构料场"],
        ],
        col_widths=[4.5, 5, 6.5],
    )

    # Section 5 - Cooperation models
    add_heading(doc, "五、建议合作模式", level=1)
    add_table(
        doc,
        ["合作模式", "具体内容", "适用阶段"],
        [
            ["技术交流", "就具体项目储运需求开展方案研讨", "拜访初期"],
            ["方案比选", "针对ML644等子项目提供储运技术方案", "1—3个月"],
            ["设备供货", "国内制造钢结构及仓体，海外指导安装", "项目招标期"],
            ["工程分包", "纳入地矿二院工程总包链条，分包储运系统", "项目执行期"],
            ["联合出海", "共同开发非洲市场储运配套项目", "中长期"],
        ],
        col_widths=[3.5, 8, 4.5],
    )

    # Section 6 - Visit strategy
    add_heading(doc, "六、拜访推进建议", level=1)

    add_heading(doc, "6.1 建议重点提及的项目", level=2)
    add_bullet(doc, "第一优先：坦桑尼亚ML644石墨矿——精矿库、湿料缓存库")
    add_bullet(doc, "第二优先：环维多利亚湖产业园——储运与设备配套")
    add_bullet(doc, "第三优先：在建/拟建金矿选厂——精矿粉储存系统")
    add_bullet(doc, "第四优先：国内嵩县等矿山——封闭堆场")

    add_heading(doc, "6.2 建议携带的资料", level=2)
    add_bullet(doc, "矿山行业解决方案（网架料场、矿粉仓）")
    add_bullet(doc, "卢旺达装配式钢板仓案例（非洲语境）")
    add_bullet(doc, "海外EPC业绩一页纸（建议附英文版）")
    add_bullet(doc, "石墨精粉/金精粉储存系统示意图")

    add_heading(doc, "6.3 拜访目标（务实设定）", level=2)
    add_table(
        doc,
        ["层级", "目标"],
        [
            ["最低目标", "建立对口部门联系，了解海外项目储运采购流程"],
            ["中等目标", "约定ML644或产业园某一子项目的技术交流"],
            ["理想目标", "进入合格供应商/分包商库，或联合编制储运方案"],
        ],
        col_widths=[3, 13],
    )

    add_heading(doc, "6.4 核心交流问题", level=2)
    add_bullet(doc, "贵院在坦桑尼亚石墨矿和金矿项目中，精矿储存和堆场封闭目前是自建还是对外招标？")
    add_bullet(doc, "海外项目储运配套的采购决策流程和归口部门是什么？")
    add_bullet(doc, "实荣能否从某一个子项目先做方案比选或技术咨询？")
    add_bullet(doc, "环维多利亚湖产业园是否有标准化储运设施采购计划？")

    # Section 7 - Risks
    add_heading(doc, "七、风险与注意事项", level=1)
    risks = [
        "国企决策链较长，一次拜访难以立竿见影，重在建立技术和采购对接关系。",
        "ML644选矿核心设备已有其他供应商，实荣宜定位储运配套分包，避免一上来谈总包。",
        "需争取与地矿二院建立直接业务联系，不完全依赖中间引荐。",
        "与龙佰等矿产合作伙伴不存在叙事冲突——地矿二院与龙佰是矿产勘查合作，与实荣是工程建设合作。",
        "双方资质可互补：地矿二院有对外承包工程资质，实荣有建安/钢结构/机电资质，可联合投标或分包。",
    ]
    for r in risks:
        add_bullet(doc, r)

    # Section 8 - Conclusion
    add_heading(doc, "八、结论", level=1)
    add_para(
        doc,
        "实荣与地矿二院的合作契合点明确，核心在于「海外矿山与选厂的散料储运配套」。"
        "尤其坦桑尼亚ML644石墨矿的精矿库/缓存库，以及环维多利亚湖产业园的储运设施，"
        "是最值得重点推进的切口。双方应定位为产业链上下游协作关系，"
        "以「贵院主导矿产开发、实荣配套储运系统」的模式，从具体子项目技术对接起步，"
        "逐步建立长期合作关系。",
    )

    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run(
        "河南实荣筒仓工程有限公司\n"
        "地址：河南省郑州市CBD商务外环路24号中国人保大厦28层\n"
        "电话：0371-63253880　邮箱：info@sronsilo.com\n"
        "网址：https://www.sronsilo.com.cn"
    )
    fr.font.size = Pt(9)
    fr.font.name = "宋体"
    fr._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    fr.font.color.rgb = RGBColor(100, 100, 100)

    return doc


if __name__ == "__main__":
    output = "/workspace/实荣与地矿二院合作契合点分析报告.docx"
    doc = build_document()
    doc.save(output)
    print(f"Generated: {output}")
